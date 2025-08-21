"""
Módulo de Prontuários - Gerencia a visualização e edição de prontuários de pacientes
"""
import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext
from tkinter import filedialog
from tkinter import font as tkfont
import win32print
from datetime import datetime
import re
import sys
try:
    # Exportação para Word (.docx)
    from docx import Document
    from docx.shared import Pt
except Exception:
    Document = None  # tratado em runtime com mensagem amigável
try:
    # Corretor ortográfico (pt)
    from spellchecker import SpellChecker
except Exception:
    SpellChecker = None
try:
    # Automação do Word (COM)
    import win32com.client as win32  # type: ignore
except Exception:
    win32 = None

from ..base_module import BaseModule
from src.controllers.cliente_controller import ClienteController
from src.controllers.prontuario_controller import ProntuarioController
from src.controllers.auth_controller import AuthController
from src.controllers.config_controller import ConfigController
from src.controllers.cadastro_controller import CadastroController
from src.controllers.agenda_controller import AgendaController
from src.utils.impressao import GerenciadorImpressao

# Filtra prints de debug específicos deste módulo, sem afetar outros prints úteis
try:
    import builtins as _bi
    def _print_filter(*args, **kwargs):
        try:
            msg = " ".join(str(a) for a in args)
            if str(msg).startswith("[DEBUG][Prontuario]"):
                return  # suprime logs de debug solicitados
        except Exception:
            pass
        return _bi.print(*args, **kwargs)
    print = _print_filter  # type: ignore
except Exception:
    pass

class ProntuarioModule(BaseModule):
    def __init__(self, parent, controller, db_connection=None):
        """
        Inicializa o módulo de prontuários.
        
        Args:
            parent: Widget pai onde este módulo será renderizado
            controller: Controlador principal da aplicação
            db_connection: Conexão com o banco de dados
        """
        super().__init__(parent, controller)
        
        # Armazena a conexão com o banco de dados
        self.db_connection = db_connection
        
        # Inicializa os controladores (prioriza os do controller principal para reutilizar a mesma conexão de BD)
        self.cliente_controller = getattr(self.controller, 'cliente_controller', None) or ClienteController()
        self.prontuario_controller = getattr(self.controller, 'prontuario_controller', None) or ProntuarioController()
        self.auth_controller = getattr(self.controller, 'auth_controller', None) or AuthController()
        self.cadastro_controller = getattr(self.controller, 'cadastro_controller', None)
        # Configurações e impressão
        self.config_controller = getattr(self.controller, 'config_controller', None) or ConfigController()
        self.impressao = GerenciadorImpressao(config_controller=self.config_controller)
        # Removido controle de fonte persistente: impressão usa padrão do sistema
        # Garante que todos os controllers usem a mesma conexão real (com .cursor)
        self._alinhar_conexoes_db()
        
        # Obtém o usuário logado (prioriza o controller principal se disponível)
        try:
            self.usuario_atual = getattr(self.controller, 'usuario', None)
        except Exception:
            self.usuario_atual = None
        if not self.usuario_atual:
            self.usuario_atual = self.auth_controller.obter_usuario_autenticado()
        # Normaliza usuário em dict e resolve médico logado
        self.usuario_dict = self._usuario_to_dict(self.usuario_atual)
        self.medico_id = None
        self._resolver_medico_logado()
        # DEBUG: identificar usuário e medico_id resolvido
        try:
            print("[DEBUG][Prontuario] usuario_dict=", self.usuario_dict, "medico_id_resolvido=", self.medico_id)
            if self.medico_id is None:
                print("[DEBUG][Prontuario] medico_id vazio (None) - usará fallback por usuario_id/nome/telefone")
        except Exception:
            pass
        
        # Variáveis de estado
        self.paciente_selecionado = None
        self.prontuario_atual = None
        self.modo_edicao = False
        # Estado do corretor ortográfico
        self._spell = None
        self._spell_after_id = None
        
        # Configura o frame principal
        self.frame.pack_propagate(False)
        
        # Constrói a interface
        self._construir_interface()
        
    def preselect_paciente_por_id(self, paciente_id):
        """Preseleciona um paciente pelo ID e atualiza a UI.
        Não exibe popups em caso de sucesso; em erro, mostra aviso discreto.
        """
        try:
            if not paciente_id:
                return
            sucesso, paciente = self.cliente_controller.buscar_cliente_por_id(int(paciente_id))
            if sucesso and paciente:
                self.paciente_selecionado = paciente
                # Reconstrói painel direito com os dados do paciente e histórico
                self._construir_painel_prontuarios()
            else:
                # Apenas alerta simples; sem alterar tema/cores
                messagebox.showwarning("Aviso", "Paciente não encontrado para pré-seleção.")
        except Exception:
            # Falha silenciosa para não interromper fluxo de navegação
            pass

    def _usuario_to_dict(self, usuario):
        """Garante um dicionário a partir do usuário autenticado (objeto ou dict)."""
        try:
            if usuario is None:
                return None
            if isinstance(usuario, dict):
                return usuario
            if hasattr(usuario, 'to_dict') and callable(getattr(usuario, 'to_dict')):
                return usuario.to_dict()
            data = {}
            for attr in ('id', 'nome', 'login', 'nivel', 'telefone'):
                if hasattr(usuario, attr):
                    data[attr] = getattr(usuario, attr)
            return data or None
        except Exception:
            return None

    def _resolver_medico_logado(self):
        """
        Define self.medico_id a partir do usuário autenticado usando o ID do usuário
        quando o nível for 'medico' ou 'funcionario'. Isso garante que os modelos de
        prontuário sejam buscados por usuário (medico_id).
        """
        # A partir de agora, salvamos e consultamos por usuario_id diretamente.
        # Não é necessário resolver/consultar medico_id.
        self.medico_id = None
        usr = getattr(self, 'usuario_dict', None) or {}
        if not isinstance(usr, dict):
            return
        user_id = usr.get('id')
        try:
            print("[DEBUG][Prontuario] usuario logado id=", user_id)
        except Exception:
            pass

    def _candidatos_identificador_medico(self):
        """Retorna possíveis identificadores do médico para compatibilidade."""
        usr = getattr(self, 'usuario_dict', None) or {}
        candidatos = []
        # id numérico do usuário (nova chave para modelos_texto)
        uid = usr.get('id')
        if uid:
            candidatos.append(uid)
        # id de médico (para outras tabelas/legados)
        if self.medico_id:
            candidatos.append(self.medico_id)
        # possíveis chaves comuns em bases existentes
        for k in ('crm', 'imp', 'codigo_medico', 'matricula', 'login', 'nome'):
            v = usr.get(k)
            if v:
                candidatos.append(v)
        # elimina duplicados preservando ordem
        seen = set()
        uniq = []
        for c in candidatos:
            if c not in seen:
                uniq.append(c)
                seen.add(c)
        # DEBUG: listar candidatos
        try:
            print("[DEBUG][Prontuario] candidatos_identificador_medico=", uniq)
        except Exception:
            pass
        return uniq

    def _obter_conexao_valida(self):
        """Retorna uma conexão que possua .cursor, priorizando a do controller principal e do cadastro_controller."""
        # 1) controller.db
        try:
            ctrl_db = getattr(self.controller, 'db', None)
            if ctrl_db and hasattr(ctrl_db, 'cursor'):
                return ctrl_db
        except Exception:
            pass
        # 2) cadastro_controller.db
        try:
            cad = getattr(self, 'cadastro_controller', None)
            cad_db = getattr(cad, 'db', None)
            if cad_db and hasattr(cad_db, 'cursor'):
                return cad_db
        except Exception:
            pass
        # 3) self.db_connection se já for conexão real
        try:
            if self.db_connection and hasattr(self.db_connection, 'cursor'):
                return self.db_connection
        except Exception:
            pass
        return None

    def _alinhar_conexoes_db(self):
        """Configura a mesma conexão válida em todos os controllers que precisarem."""
        conn = self._obter_conexao_valida()
        try:
            print("[DEBUG][Prontuario] alinhando conexões: conexao_valida_existe=", bool(conn))
        except Exception:
            pass
        if not conn:
            return
        # Atribui nos controllers com set_db_connection ou atributo .db
        for ctrl in (self.cliente_controller, self.prontuario_controller, self.auth_controller, self.cadastro_controller):
            if not ctrl:
                continue
            try:
                if hasattr(ctrl, 'set_db_connection') and callable(getattr(ctrl, 'set_db_connection')):
                    ctrl.set_db_connection(conn)
                elif hasattr(ctrl, 'db'):
                    setattr(ctrl, 'db', conn)
            except Exception:
                continue

    def _listar_modelos_compat(self):
        """Tenta listar modelos por múltiplos identificadores do médico para compatibilidade."""
        erros = []
        # tentar pelas APIs conhecidas, variando o identificador
        for ident in self._candidatos_identificador_medico():
            try:
                print(f"[DEBUG][Prontuario] tentando listar modelos com ident={ident}")
            except Exception:
                pass
            # tentar listar_modelos_texto somente quando ident for usuario_id numérico
            try:
                ident_uid = None
                if isinstance(ident, int):
                    ident_uid = ident
                else:
                    # strings numéricas
                    try:
                        ident_uid = int(str(ident)) if str(ident).isdigit() else None
                    except Exception:
                        ident_uid = None
                modelos = []
                if ident_uid is not None:
                    modelos = self.prontuario_controller.listar_modelos_texto(ident_uid)
                if modelos:
                    try:
                        print(f"[DEBUG][Prontuario] listar_modelos_texto encontrou {len(modelos)} modelos para usuario_id={ident_uid}")
                    except Exception:
                        pass
                    return modelos
            except Exception as e:
                erros.append(str(e))
                try:
                    print(f"[DEBUG][Prontuario] erro listar_modelos_texto usuario_id={ident_uid}: {e}")
                except Exception:
                    pass
            # tentar buscar_modelos_texto
            try:
                modelos = getattr(self.prontuario_controller, 'buscar_modelos_texto', lambda *_: [])(ident)
                if modelos:
                    try:
                        print(f"[DEBUG][Prontuario] buscar_modelos_texto encontrou {len(modelos)} modelos para ident={ident}")
                    except Exception:
                        pass
                    return modelos
            except Exception as e:
                erros.append(str(e))
                try:
                    print(f"[DEBUG][Prontuario] erro buscar_modelos_texto ident={ident}: {e}")
                except Exception:
                    pass
        # se nada encontrado, retorna lista vazia; erros podem ser mostrados em chamadas
        return []

    def _construir_interface(self):
        """Constrói a interface do módulo de prontuários."""
        # Frame principal dividido em duas partes
        self.painel_esquerdo = tk.Frame(self.frame, bg='#f0f2f5', width=300)
        self.painel_esquerdo.pack(side=tk.LEFT, fill=tk.Y, padx=10, pady=10)
        self.painel_esquerdo.pack_propagate(False)
        
        self.painel_direito = tk.Frame(self.frame, bg='#f0f2f5')
        self.painel_direito.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Constrói o painel de busca de pacientes
        self._construir_painel_busca()
        
        # Constrói o painel de prontuários (inicialmente vazio)
        self._construir_painel_prontuarios()
    
    def _construir_painel_busca(self):
        """Constrói o painel de busca de pacientes."""
        # Frame para o título
        titulo_frame = tk.Frame(self.painel_esquerdo, bg='#f0f2f5')
        titulo_frame.pack(fill=tk.X, pady=(0, 10))
        
        tk.Label(
            titulo_frame,
            text="Buscar Paciente",
            font=('Arial', 14, 'bold'),
            bg='#f0f2f5',
            fg='#333333'
        ).pack(side=tk.LEFT)
        
        # Frame para a busca
        busca_frame = tk.Frame(self.painel_esquerdo, bg='#f0f2f5')
        busca_frame.pack(fill=tk.X, pady=(0, 10))
        
        # Campo de busca
        self.busca_var = tk.StringVar()
        self.busca_entry = ttk.Entry(
            busca_frame,
            textvariable=self.busca_var,
            width=25
        )
        self.busca_entry.pack(side=tk.LEFT, padx=(0, 5))
        self.busca_entry.bind("<Return>", self._buscar_paciente)
        # Busca dinâmica ao digitar (debounce)
        self._busca_after_id = None
        self.busca_entry.bind("<KeyRelease>", self._on_change_busca)
        
        # Botão de busca
        tk.Button(
            busca_frame,
            text="Buscar",
            command=self._buscar_paciente,
            bg='#4a6fa5',      # Cor de fundo do botão
            fg='#ffffff',      # Cor do texto
            activebackground='#3b5a7f',  # Cor quando o botão é pressionado
            activeforeground='#ffffff',  # Cor do texto quando pressionado
            relief=tk.FLAT,  # Estilo do relevo do botão
            borderwidth=2,     # Largura da borda
            font=('Arial', 10, 'bold'),  # Fonte em negrito
            cursor='hand2'     # Cursor de mão ao passar por cima
        ).pack(side=tk.LEFT, padx=(0, 5))  # Adiciona um pequeno espaço à direita

        # Frame para a lista de resultados
        resultados_frame = tk.Frame(self.painel_esquerdo, bg='#f0f2f5')
        resultados_frame.pack(fill=tk.BOTH, expand=True)
                
        # Lista de resultados com scrollbar
        self.resultados_tree = ttk.Treeview(
            resultados_frame,
            columns=("id", "nome"),
            show="headings",
            height=15
        )
        
        # Configuração das colunas
        self.resultados_tree.heading("id", text="ID")
        self.resultados_tree.heading("nome", text="Nome")
        
        self.resultados_tree.column("id", width=50)
        self.resultados_tree.column("nome", width=220)
        
        # Scrollbar
        scrollbar = ttk.Scrollbar(
            resultados_frame,
            orient="vertical",
            command=self.resultados_tree.yview
        )
        self.resultados_tree.configure(yscrollcommand=scrollbar.set)
        
        # Posicionamento
        self.resultados_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Evento de seleção
        self.resultados_tree.bind("<<TreeviewSelect>>", self._selecionar_paciente)
    
    def _construir_painel_prontuarios(self):
        """Constrói o painel de prontuários."""
        # Limpa o painel direito se já existir algum conteúdo
        for widget in self.painel_direito.winfo_children():
            widget.destroy()
            
        # Frame principal do painel de prontuários
        conteudo_frame = tk.Frame(self.painel_direito, bg='#f0f2f5')
        conteudo_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Se não houver paciente selecionado, exibe mensagem
        if not self.paciente_selecionado:
            tk.Label(
                conteudo_frame,
                text="Selecione um paciente para visualizar ou criar prontuários",
                font=('Arial', 12),
                bg='#f0f2f5',
                fg='#666666',
                pady=50
            ).pack(expand=True)
            return
            
        # Nenhum botão de ação no topo deste painel
        
        # Frame para a área esquerda (carrossel) e conteúdo à direita (editor)
        prontuario_container = tk.Frame(conteudo_frame, bg='#f0f2f5')
        prontuario_container.pack(fill=tk.BOTH, expand=True)
        
        # Frame para o carrossel de prontuários (esquerda)
        carrossel_frame = tk.Frame(prontuario_container, bg='#f0f2f5', width=380)
        carrossel_frame.pack(side=tk.LEFT, fill=tk.BOTH, padx=(0, 10))
        carrossel_frame.pack_propagate(False)
        
        # Frame para o conteúdo do prontuário (editor à direita)
        self.prontuario_content_frame = tk.Frame(prontuario_container, bg='#ffffff', relief=tk.SUNKEN, bd=1)
        self.prontuario_content_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        # Renderiza editor do novo prontuário
        self._render_editor_novo_prontuario(self.prontuario_content_frame)
        
        # Renderiza o carrossel do histórico do paciente
        self._render_carrossel_prontuarios(carrossel_frame)
        # Carrega os prontuários do paciente no carrossel
        self._carregar_prontuarios()
    
    def _buscar_paciente(self, event=None):
        """Busca pacientes com base no texto inserido. Pode ser chamada por botão ou tecla Enter."""
        termo_busca = self.busca_var.get().strip()
        
        if not termo_busca:
            messagebox.showinfo("Aviso", "Digite um termo para busca")
            return
        
        # Limpa a lista de resultados
        for item in self.resultados_tree.get_children():
            self.resultados_tree.delete(item)
        
        # Verifica se o termo é numérico (possível ID ou telefone)
        if termo_busca.isdigit():
            # Busca por ID
            sucesso, paciente = self.cliente_controller.buscar_cliente_por_id(int(termo_busca))
            if sucesso and paciente:
                self.resultados_tree.insert("", "end", values=(
                    paciente["id"],
                    paciente["nome"]
                ))
            
            # Busca por telefone
            pacientes = self.cliente_controller.buscar_cliente_por_telefone(termo_busca)
            for paciente in pacientes:
                # Evita duplicatas se já encontrou por ID
                if not sucesso or paciente["id"] != int(termo_busca):
                    self.resultados_tree.insert("", "end", values=(
                        paciente["id"],
                        paciente["nome"]
                    ))
        else:
            # Busca por nome
            pacientes = self.cliente_controller.buscar_cliente_por_nome(termo_busca)
            for paciente in pacientes:
                self.resultados_tree.insert("", "end", values=(
                    paciente["id"],
                    paciente["nome"]
                ))
        
        # Verifica se encontrou resultados
        if not self.resultados_tree.get_children():
            messagebox.showinfo("Resultado", "Nenhum paciente encontrado")
        
        # Se foi chamado via evento de teclado, evita propagação
        return "break"
    
    def _selecionar_paciente(self, event):
        """Seleciona um paciente da lista de resultados."""
        selecao = self.resultados_tree.selection()
        if not selecao:
            return
        
        # Obtém o ID do paciente selecionado
        item = self.resultados_tree.item(selecao[0])
        paciente_id = item["values"][0]
        
        # Busca os dados completos do paciente
        sucesso, paciente = self.cliente_controller.buscar_cliente_por_id(paciente_id)
        if sucesso and paciente:
            self.paciente_selecionado = paciente
            # Reconstrói o painel de prontuários com o novo paciente
            self._construir_painel_prontuarios()
        else:
            messagebox.showerror("Erro", "Não foi possível carregar os dados do paciente")
    
    def _carregar_prontuarios(self):
        """Carrega os prontuários do paciente selecionado no carrossel."""
        if not self.paciente_selecionado:
            return
        prontuarios = self.prontuario_controller.buscar_prontuarios_paciente(self.paciente_selecionado["id"])
        # Ordena localmente: mais recentes primeiro (data desc, depois id desc)
        try:
            lista = list(prontuarios or [])
            lista.sort(key=lambda p: (
                (p.get('data') or ''),
                (p.get('id') or 0)
            ), reverse=True)
            self.prontuarios_lista = lista
        except Exception:
            # Fallback sem ordenação caso algo falhe
            self.prontuarios_lista = prontuarios or []
        self.carousel_index = 0 if self.prontuarios_lista else -1
        self._atualizar_carrossel()

    def _render_carrossel_prontuarios(self, parent):
        """Monta a UI do carrossel do histórico do paciente."""
        # Título
        tk.Label(
            parent,
            text=f"Histórico de {self.paciente_selecionado['nome']}",
            font=('Arial', 11, 'bold'),
            bg='#f0f2f5',
            fg='#333333'
        ).pack(fill=tk.X, pady=(0, 5))

        # Navegação
        nav = tk.Frame(parent, bg='#f0f2f5')
        nav.pack(fill=tk.X, pady=(0, 5))
        self.btn_prev = tk.Button(
            nav, text='◀', width=3, command=self._carrossel_prev,
            bg='#4a6fa5', fg='#ffffff', activebackground='#3b5a7f', activeforeground='#ffffff',
            relief=tk.FLAT, cursor='hand2'
        )
        self.btn_prev.pack(side=tk.LEFT)
        self.carousel_pos_label = tk.Label(nav, text='0/0', bg='#f0f2f5')
        self.carousel_pos_label.pack(side=tk.LEFT, padx=8)
        self.btn_next = tk.Button(
            nav, text='▶', width=3, command=self._carrossel_next,
            bg='#4a6fa5', fg='#ffffff', activebackground='#3b5a7f', activeforeground='#ffffff',
            relief=tk.FLAT, cursor='hand2'
        )
        self.btn_next.pack(side=tk.LEFT)
        # Botão para editar/salvar o item atual do carrossel no Word
        try:
            self.btn_edit_save = tk.Button(
                nav, text='Editar', width=8, command=self._carrossel_editar,
                bg='#4a6fa5', fg='#ffffff', activebackground='#3b5a7f', activeforeground='#ffffff',
                relief=tk.FLAT, cursor='hand2'
            )
            self.btn_edit_save.pack(side=tk.LEFT, padx=8)
        except Exception:
            pass
        # Botão Excluir ao lado do Editar
        self.btn_excluir = tk.Button(
            nav, text='Excluir', width=8, command=self._carrossel_excluir,
            bg='#dc3545', fg='#ffffff', activebackground='#c82333', activeforeground='#ffffff',
            relief=tk.FLAT, cursor='hand2'
        )
        self.btn_excluir.pack(side=tk.LEFT, padx=(0, 8))

        # Preview
        preview_frame = tk.Frame(parent, bg='#f0f2f5')
        preview_frame.pack(fill=tk.BOTH, expand=True)
        self.carousel_preview = scrolledtext.ScrolledText(
            preview_frame,
            wrap=tk.WORD,
            font=('Arial', 10),
            state=tk.DISABLED,
            height=18
        )
        self.carousel_preview.pack(fill=tk.BOTH, expand=True)

    def _atualizar_carrossel(self):
        """Atualiza o preview e estado dos botões do carrossel."""
        total = len(getattr(self, 'prontuarios_lista', []) or [])
        idx = getattr(self, 'carousel_index', -1)
        # Atualiza label de posição
        self.carousel_pos_label.config(text=f"{(idx+1) if total and idx>=0 else 0}/{total}")
        # Habilita/desabilita botões
        state_prev = tk.NORMAL if idx > 0 else tk.DISABLED
        state_next = tk.NORMAL if (total and idx < total - 1) else tk.DISABLED
        self.btn_prev.config(state=state_prev)
        self.btn_next.config(state=state_next)

        # Atualiza conteúdo
        self.carousel_preview.config(state=tk.NORMAL)
        self.carousel_preview.delete(1.0, tk.END)
        if total and 0 <= idx < total:
            item = self.prontuarios_lista[idx]
            # Mostra apenas o conteúdo salvo, sem inserir título ou cabeçalhos extras
            self.carousel_preview.insert(tk.END, item.get('conteudo', ''))
        else:
            self.carousel_preview.insert(tk.END, 'Nenhum prontuário encontrado.')
        self.carousel_preview.config(state=tk.DISABLED)

    def _carrossel_prev(self):
        if getattr(self, 'carousel_index', -1) > 0:
            self.carousel_index -= 1
            self._atualizar_carrossel()

    def _carrossel_next(self):
        total = len(getattr(self, 'prontuarios_lista', []) or [])
        if total and getattr(self, 'carousel_index', -1) < total - 1:
            self.carousel_index += 1
            self._atualizar_carrossel()

    def _carrossel_editar(self):
        """Abre o Word com o conteúdo do item atual do carrossel para edição e troca o botão para 'Salvar'."""
        try:
            total = len(getattr(self, 'prontuarios_lista', []) or [])
            idx = getattr(self, 'carousel_index', -1)
            if not (total and 0 <= idx < total):
                return
            item = self.prontuarios_lista[idx]
            conteudo = item.get('conteudo') or ''
            if win32 is None:
                messagebox.showwarning('Aviso', 'Automação do Word não está disponível nesta máquina.')
                return
            # Instância do Word
            try:
                word_app = getattr(self, '_word_app_carousel', None) or win32.gencache.EnsureDispatch('Word.Application')
            except Exception:
                word_app = getattr(self, '_word_app_carousel', None) or win32.Dispatch('Word.Application')
            word_app.Visible = True
            self._word_app_carousel = word_app
            # Fecha documento anterior deste fluxo, se houver
            try:
                if getattr(self, '_word_doc_carousel', None) is not None:
                    self._word_doc_carousel.Close(SaveChanges=0)
            except Exception:
                pass
            # Cria documento e injeta o texto (\n -> \r)
            doc = word_app.Documents.Add()
            self._word_doc_carousel = doc
            doc.Range(0, 0).Text = (conteudo or '').replace('\n', '\r')
            # Troca o botão para 'Salvar' (verde)
            if hasattr(self, 'btn_edit_save'):
                self.btn_edit_save.config(
                    text='Salvar', command=self._carrossel_salvar,
                    bg='#28a745', fg='#ffffff', activebackground='#218838', activeforeground='#ffffff'
                )
        except Exception as e:
            messagebox.showerror('Erro', f'Falha ao abrir no Word: {e}')

    def _carrossel_salvar(self):
        """Salva o conteúdo do Word no mesmo registro e atualiza o preview, mantendo posição/quantidade."""
        try:
            total = len(getattr(self, 'prontuarios_lista', []) or [])
            idx = getattr(self, 'carousel_index', -1)
            if not (total and 0 <= idx < total):
                return
            item = self.prontuarios_lista[idx]
            prontuario_id = item.get('id')
            if not prontuario_id:
                messagebox.showwarning('Aviso', 'Registro do prontuário não encontrado para salvar.')
                return
            # Lê texto do Word
            raw = ''
            try:
                if getattr(self, '_word_doc_carousel', None) is not None:
                    raw = self._word_doc_carousel.Content.Text or ''
            except Exception:
                raw = ''
            texto = (raw or '').replace('\r', '\n').rstrip('\n')
            # Atualiza no banco
            ok, msg = self.prontuario_controller.atualizar_prontuario(prontuario_id, {'conteudo': texto})
            if not ok:
                messagebox.showerror('Erro', msg)
                return
            # Atualiza em memória e preview
            self.prontuarios_lista[idx]['conteudo'] = texto
            self._atualizar_carrossel()
            # Fecha doc do Word (mantém app para reuso)
            try:
                if getattr(self, '_word_doc_carousel', None) is not None:
                    self._word_doc_carousel.Close(SaveChanges=0)
            except Exception:
                pass
            self._word_doc_carousel = None
            # Volta botão para 'Editar' (azul)
            if hasattr(self, 'btn_edit_save'):
                self.btn_edit_save.config(
                    text='Editar', command=self._carrossel_editar,
                    bg='#4a6fa5', fg='#ffffff', activebackground='#3b5a7f', activeforeground='#ffffff'
                )
           
        except Exception as e:
            messagebox.showerror('Erro', f'Falha ao salvar alterações: {e}')

    def _carrossel_excluir(self):
        """Exclui o prontuário atual do carrossel após confirmação, atualizando a lista e o preview."""
        try:
            total = len(getattr(self, 'prontuarios_lista', []) or [])
            idx = getattr(self, 'carousel_index', -1)
            if not (total and 0 <= idx < total):
                return
            item = self.prontuarios_lista[idx]
            prontuario_id = item.get('id')
            if not prontuario_id:
                messagebox.showwarning('Aviso', 'Registro do prontuário não encontrado para excluir.')
                return
            # Confirmação
            if not messagebox.askyesno('Excluir', 'Confirma a exclusão deste prontuário? Esta ação não pode ser desfeita.'):
                return
            # Se houver documento Word aberto no fluxo do carrossel, fecha sem salvar
            try:
                if getattr(self, '_word_doc_carousel', None) is not None:
                    self._word_doc_carousel.Close(SaveChanges=0)
            except Exception:
                pass
            self._word_doc_carousel = None
            # Chama controller para excluir
            ok, msg = self.prontuario_controller.excluir_prontuario(prontuario_id)
            if not ok:
                messagebox.showerror('Erro', msg)
                return
            # Remove da lista e ajusta índice
            try:
                del self.prontuarios_lista[idx]
            except Exception:
                # fallback: reconstruir
                self.prontuarios_lista = [p for p in self.prontuarios_lista if p.get('id') != prontuario_id]
            novo_total = len(self.prontuarios_lista)
            if novo_total == 0:
                self.carousel_index = -1
            else:
                if idx >= novo_total:
                    idx = novo_total - 1
                self.carousel_index = idx
            # Restaura estado do botão editar
            if hasattr(self, 'btn_edit_save'):
                self.btn_edit_save.config(
                    text='Editar', command=self._carrossel_editar,
                    bg='#4a6fa5', fg='#ffffff', activebackground='#3b5a7f', activeforeground='#ffffff'
                )
            # Atualiza UI
            self._atualizar_carrossel()
        except Exception as e:
            messagebox.showerror('Erro', f'Falha ao excluir prontuário: {e}')

    def _exibir_prontuario(self):
        """Exibe o conteúdo do prontuário selecionado."""
        # Limpa o frame de conteúdo
        for widget in self.prontuario_content_frame.winfo_children():
            widget.destroy()
        
        if not self.prontuario_atual:
            return
        
        # Habilita a edição do texto para inserir o conteúdo
        self.texto_prontuario = scrolledtext.ScrolledText(
            self.prontuario_content_frame,
            wrap=tk.WORD,
            font=('Arial', 11),
            height=20
        )
        self.texto_prontuario.pack(fill=tk.BOTH, expand=True)
        
        # Limpa o conteúdo atual
        self.texto_prontuario.delete(1.0, tk.END)
        
        # Insere apenas o conteúdo do prontuário (sem cabeçalho extra)
        conteudo = self.prontuario_atual.get("conteudo", "")
        self.texto_prontuario.insert(tk.END, conteudo, "conteudo")
        
        # Configura as tags de formatação
        self.texto_prontuario.tag_configure("titulo", font=("Arial", 12, "bold"))
        self.texto_prontuario.tag_configure("info", font=("Arial", 10, "italic"))
        self.texto_prontuario.tag_configure("conteudo", font=("Arial", 10))
        
        # Desabilita a edição do texto (modo somente leitura)
        self.texto_prontuario.config(state=tk.DISABLED)
        
        # Rola para o início do texto
        self.texto_prontuario.see("1.0")

    def _render_editor_novo_prontuario(self, parent):
        """Renderiza o editor do novo prontuário à direita."""
        # Cabeçalho do editor
        header = tk.Frame(parent, bg='#ffffff')
        header.pack(fill=tk.X)
        tk.Label(
            header,
            text="Novo Prontuário",
            bg='#ffffff',
            fg='#333333',
            font=('Arial', 12, 'bold')
        ).pack(side=tk.LEFT, padx=8, pady=8)

        # Barra de botões
        botoes = tk.Frame(parent, bg='#ffffff')
        botoes.pack(fill=tk.X, padx=8, pady=(0, 6))

        # Botões (mantém tema/cores existentes)
        tk.Button(
            botoes, text="Salvar", command=self._salvar_novo_prontuario,
            bg="#28a745", fg="#ffffff", activebackground="#218838", activeforeground="#ffffff",
            relief=tk.FLAT, cursor='hand2'
        ).pack(side=tk.RIGHT, padx=4)

        tk.Button(
            botoes, text="Prontuarios", command=self._selecionar_modelo,
            bg="#4a6fa5", fg="white", relief=tk.FLAT, cursor='hand2'
        ).pack(side=tk.RIGHT, padx=4)

        tk.Button(
            botoes, text="Receitas", command=self._selecionar_receita,
            bg="#4a6fa5", fg="white", relief=tk.FLAT, cursor='hand2'
        ).pack(side=tk.RIGHT, padx=4)

        def _limpar():
            try:
                self.editor_texto.delete('1.0', tk.END)
                self._preencher_cabecalho_novo_prontuario()
            except Exception:
                pass
        tk.Button(
            botoes, text="Limpar", command=_limpar,
            bg="#dc3545", fg="white", relief=tk.FLAT, cursor='hand2'
        ).pack(side=tk.RIGHT, padx=4)

        # Imprimir: mantém fluxo existente; se método não existir, mostra aviso discreto
        def _imprimir_dispatch():
            try:
                func = getattr(self, '_imprimir_prontuario', None)
                if callable(func):
                    return func()
                messagebox.showinfo('Impressão', 'Função de impressão indisponível nesta versão.')
            except Exception:
                pass
        tk.Button(
            botoes, text="Imprimir", command=_imprimir_dispatch,
            bg="#17a2b8", fg="white", relief=tk.FLAT, cursor='hand2'
        ).pack(side=tk.RIGHT, padx=4)

        # Novo: Exportar Word (.docx)
        self._btn_word = tk.Button(
            botoes, text="Editar", command=self._toggle_word_roundtrip,
            bg="#4a6fa5", fg="white", relief=tk.FLAT, cursor='hand2'
        )
        self._btn_word.pack(side=tk.RIGHT, padx=4)

        # Editor de texto
        self.editor_texto = scrolledtext.ScrolledText(
            parent,
            wrap=tk.WORD,
            font=('Arial', 11),
            height=22
        )
        self.editor_texto.pack(fill=tk.BOTH, expand=True, padx=8, pady=(0, 8))

        # Estado de edição ativo
        self.modo_edicao = True
        # Prefill de cabeçalho com paciente + data
        self._preencher_cabecalho_novo_prontuario()

   

    def _imprimir_prontuario(self):
        """Imprime o texto atual do editor em A4 usando o GerenciadorImpressao.
        Não altera o conteúdo: envia exatamente o que está no editor (diretivas <<font:>> e alinhamento são suportados pelo motor de impressão).
        """
        try:
            txt = getattr(self, 'editor_texto', None)
            if not txt:
                messagebox.showwarning('Impressão', 'Editor não está disponível para impressão.')
                return
            conteudo = txt.get('1.0', 'end-1c')
            if not conteudo.strip():
                messagebox.showwarning('Impressão', 'Não há conteúdo para imprimir.')
                return

            # Seleciona impressora (usa diálogo existente)
            alvo = self._dialogo_escolher_sala()
            if not alvo:
                return  # cancelado

            # Imprime como texto A4 simples (usa método público que encaminha ao A4)
            ok = False
            try:
                ok = self.impressao.imprimir_receita_texto(conteudo, impressora=alvo)
            except Exception:
                ok = False
            if not ok:
                messagebox.showerror('Impressão', 'Falha ao enviar para a impressora. Verifique a impressora configurada.')
        except Exception as e:
            messagebox.showerror('Impressão', f'Erro ao imprimir: {e}')

    def _exportar_para_word(self):
        """Exporta o conteúdo do editor para um arquivo Word (.docx)."""
        try:
            # Coleta dados de cabeçalho
            nome_paciente = ''
            try:
                if getattr(self, 'paciente_selecionado', None):
                    nome_paciente = self.paciente_selecionado.get('nome') or ''
            except Exception:
                nome_paciente = ''
            medico_info = None
            try:
                medico_info = self._buscar_medico_por_usuario()
            except Exception:
                medico_info = None
            nome_medico = ''
            try:
                usr = getattr(self, 'usuario_dict', None) or {}
                nome_medico = (medico_info.get('nome') if medico_info else None) or (usr.get('nome') or '')
            except Exception:
                pass
            data_hoje = datetime.now().strftime('%d/%m/%Y')

            # Obtém conteúdo do editor (texto simples)
            txt_widget = getattr(self, 'editor_texto', None)
            if not txt_widget:
                messagebox.showwarning('Aviso', 'Editor não está disponível para exportação.')
                return
            conteudo = txt_widget.get('1.0', 'end-1c')

            # Detecta diretiva de fonte global (<<font:N>>) e remove-a do conteúdo
            base_pt = 11
            linhas = conteudo.replace('\r\n', '\n').replace('\r', '\n').split('\n')
            for i, ln in enumerate(list(linhas)):
                s = ln.strip()
                if s.lower().startswith('<<font:') and s.endswith('>>'):
                    try:
                        v = int(s[7:-2])
                        if 6 <= v <= 20:
                            base_pt = v
                    except Exception:
                        pass
                    # Remove a linha da diretiva do texto exportado
                    try:
                        del linhas[i]
                    except Exception:
                        pass
                    break

            # Remove cabeçalho duplicado no início do texto (Paciente/Médico/Data)
            try:
                def _strip_initial_header(ls):
                    i = 0
                    seen = {'paciente': False, 'medico': False, 'data': False}
                    while i < len(ls):
                        s = (ls[i] or '').strip()
                        if not s:
                            i += 1
                            continue
                        low = s.lower()
                        if low.startswith('paciente:') or low.startswith('paciente :'):
                            seen['paciente'] = True
                            i += 1
                            continue
                        if low.startswith('médico:') or low.startswith('medico:') or low.startswith('médico :') or low.startswith('medico :'):
                            seen['medico'] = True
                            i += 1
                            continue
                        if low.startswith('data:') or low.startswith('data :'):
                            seen['data'] = True
                            i += 1
                            continue
                        break
                    if seen['paciente'] or seen['medico'] or seen['data']:
                        # remove linhas consideradas cabeçalho e espaços em branco seguintes
                        while i < len(ls) and not (ls[i] or '').strip():
                            i += 1
                        return ls[i:]
                    return ls
                linhas = _strip_initial_header(linhas)
            except Exception:
                pass

            # Tenta abrir diretamente no Word via COM (preferido)
            if win32 is not None:
                try:
                    word_app = getattr(self, '_word_app_export', None) or win32.gencache.EnsureDispatch('Word.Application')
                    word_app.Visible = True
                    self._word_app_export = word_app
                    doc = word_app.Documents.Add()
                    sel = word_app.Selection

                    # Cabeçalho: Paciente, Médico, Data (Arial 11, bold)
                    header_lines = []
                    if nome_paciente:
                        header_lines.append(f"Paciente: {nome_paciente}")
                    if nome_medico:
                        header_lines.append(f"Médico: {nome_medico}")
                    header_lines.append(f"Data: {data_hoje}")

                    for hl in header_lines:
                        try:
                            sel.Font.Name = 'Arial'
                        except Exception:
                            pass
                        sel.Font.Size = 11
                        sel.Font.Bold = True
                        sel.TypeText(hl)
                        sel.TypeParagraph()
                    sel.TypeParagraph()  # linha em branco
                    sel.Font.Bold = False

                    # Corpo: respeita <<center>> e <<right>>; fonte base_pt
                    # WdParagraphAlignment: Left=0, Center=1, Right=2, Justify=3
                    for linha in linhas:
                        raw = linha.lstrip()
                        align = 0
                        text = raw
                        low = raw.lower()
                        if low.startswith('<<center>>'):
                            align = 1
                            text = raw[10:].lstrip()
                        elif low.startswith('<<right>>'):
                            align = 2
                            text = raw[9:].lstrip()
                        sel.ParagraphFormat.Alignment = align
                        try:
                            sel.Font.Name = 'Arial'
                        except Exception:
                            pass
                        sel.Font.Size = base_pt
                        sel.TypeText(text)
                        sel.TypeParagraph()

                    # Guarda referência do documento e habilita retorno
                    try:
                        self._word_doc_export = doc
                        self._word_return_ready = True
                        if hasattr(self, '_btn_word') and self._btn_word.winfo_exists():
                            self._btn_word.config(text='Retornar')
                    except Exception:
                        pass
                    return
                except Exception:
                    # Se falhar COM, cai no fluxo de salvar .docx
                    pass

            # Fallback: fluxo atual de salvar .docx
            # Verifica dependência
            if Document is None:
                messagebox.showwarning('Recurso indisponível', 'Biblioteca python-docx não está instalada. Instale para usar a exportação para Word (.docx).')
                return

            # Dialogo de salvar
            sugestao_nome = f"Prontuario_{nome_paciente.replace(' ', '_')}_{datetime.now().strftime('%Y-%m-%d')}.docx" if nome_paciente else f"Prontuario_{datetime.now().strftime('%Y-%m-%d')}.docx"
            caminho = filedialog.asksaveasfilename(
                title='Salvar como',
                defaultextension='.docx',
                filetypes=[('Documento do Word', '*.docx')],
                initialfile=sugestao_nome
            )
            if not caminho:
                return

            # Monta documento
            document = Document()

            # Cabeçalho: Paciente, Médico, Data
            header_lines = []
            if nome_paciente:
                header_lines.append(f"Paciente: {nome_paciente}")
            if nome_medico:
                header_lines.append(f"Médico: {nome_medico}")
            header_lines.append(f"Data: {data_hoje}")

            for linha in header_lines:
                run = document.add_paragraph().add_run(linha)
                f = run.font
                try:
                    f.name = 'Arial'
                except Exception:
                    pass
                f.size = Pt(11)
                f.bold = True

            document.add_paragraph('')  # linha em branco

            # Conteúdo: mantém texto simples, removendo marcadores de alinhamento
            # Remove marcadores reconhecidos do início de cada linha: <<center>>, <<right>>
            for linha in linhas:
                low = linha.lstrip().lower()
                if low.startswith('<<center>>'):
                    linha = linha.lstrip()[10:].lstrip()
                elif low.startswith('<<right>>'):
                    linha = linha.lstrip()[9:].lstrip()
                run = document.add_paragraph().add_run(linha)
                f = run.font
                try:
                    f.name = 'Arial'
                except Exception:
                    pass
                f.size = Pt(base_pt)

            # Salva
            document.save(caminho)
            messagebox.showinfo('Exportação', f'Prontuário exportado com sucesso para:\n{caminho}')
            # Marca caminho para retorno e alterna botão
            try:
                self._last_export_docx_path = caminho
                self._word_return_ready = True
                if hasattr(self, '_btn_word') and self._btn_word.winfo_exists():
                    self._btn_word.config(text='Salvar')
            except Exception:
                pass
        except Exception as e:
            messagebox.showerror('Erro ao exportar', f'{e}')

    def _toggle_word_roundtrip(self):
        """Alterna entre exportar para o Word e retornar o conteúdo editado."""
        try:
            if getattr(self, '_word_return_ready', False):
                # Retornar conteúdo para o editor
                self._importar_do_word()
                # Reset estado e botão
                self._word_return_ready = False
                if hasattr(self, '_btn_word') and self._btn_word.winfo_exists():
                    self._btn_word.config(text='Editar')
            else:
                # Exportar
                self._exportar_para_word()
                # Se a exportação preparou retorno, o texto do botão já foi ajustado lá
        except Exception:
            pass

    def _importar_do_word(self):
        """Importa o conteúdo do documento Word aberto/salvo de volta para o editor_texto.
        Prioriza COM (documento aberto). Caso contrário, tenta último .docx exportado.
        """
        try:
            txt_widget = getattr(self, 'editor_texto', None)
            if not txt_widget:
                messagebox.showwarning('Aviso', 'Editor não está disponível para importar o conteúdo.')
                return

            conteudo_retorno = None

            # 1) COM: documento aberto no Word
            try:
                if win32 is not None and getattr(self, '_word_app_export', None) is not None:
                    doc = getattr(self, '_word_doc_export', None)
                    if doc is not None:
                        # Captura todo o texto do documento
                        raw = doc.Content.Text
                        # Normaliza quebras de linha
                        conteudo_retorno = (raw or '').replace('\r\x07', '\n').replace('\r', '\n')
            except Exception:
                conteudo_retorno = None

            # 2) Fallback: último arquivo .docx salvo
            if conteudo_retorno is None:
                try:
                    caminho = getattr(self, '_last_export_docx_path', None)
                    if caminho and Document is not None:
                        docx = Document(caminho)
                        # Concatena parágrafos com \n
                        conteudo_retorno = "\n".join(p.text or '' for p in docx.paragraphs)
                except Exception:
                    conteudo_retorno = None

            if not conteudo_retorno:
                messagebox.showinfo('Retornar do Word', 'Não foi possível obter o conteúdo do Word. Salve o documento e tente novamente.')
                return

            # Atualiza o editor com o conteúdo retornado
            try:
                txt_widget.delete('1.0', tk.END)
                txt_widget.insert(tk.END, conteudo_retorno)
                # Reaplica correção ortográfica, se habilitada
                try:
                    self._verificar_ortografia()
                except Exception:
                    pass
            except Exception as e:
                messagebox.showerror('Erro', f'Falha ao atualizar conteúdo do editor: {e}')
        except Exception:
            pass

    def _dialogo_escolher_sala(self):
        """Dialoga a escolha de Sala de Impressão (1..5) e retorna o nome da impressora configurada ou None.
        - Janela centralizada e com tamanho adequado.
        - Sem ttk.Button (usa tk.Button no padrão do app).
        """
        # Carrega mapeamento de salas -> impressoras
        cfg = {}
        try:
            if self.config_controller and hasattr(self.config_controller, 'carregar_config_impressoras'):
                cfg = self.config_controller.carregar_config_impressoras() or {}
        except Exception:
            cfg = {}

        itens = []
        for i in range(1, 6):
            nome_imp = cfg.get(f'impressora {i}') or ''
            label = f"Impressora {i} — {nome_imp if nome_imp else 'Não configurada'}"
            itens.append((label, nome_imp))

        sel = {'valor': None}
        dlg = tk.Toplevel(self.parent)
        dlg.title("Selecionar Impressora")
        dlg.configure(bg='#ffffff')
        dlg.transient(self.parent)
        dlg.grab_set()
        # Dimensão amigável
        W, H = 560, 420
        try:
            sw = dlg.winfo_screenwidth(); sh = dlg.winfo_screenheight()
            x = (sw - W) // 2; y = (sh - H) // 2
            dlg.geometry(f"{W}x{H}+{x}+{y}")
        except Exception:
            dlg.geometry(f"{W}x{H}")
        dlg.minsize(520, 360)

        tk.Label(dlg, text="Escolha a impressora:", bg='#ffffff', fg='#333333', font=('Arial', 11, 'bold')).pack(padx=12, pady=(12, 8))

        # Área da lista com scrollbar
        cont = tk.Frame(dlg, bg='#ffffff')
        cont.pack(fill=tk.BOTH, expand=True, padx=12)
        sb = tk.Scrollbar(cont, orient=tk.VERTICAL)
        lb = tk.Listbox(cont, height=10, yscrollcommand=sb.set)
        sb.config(command=lb.yview)
        lb.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        for label, _ in itens:
            lb.insert(tk.END, label)
        if itens:
            lb.selection_set(0)
            lb.see(0)

        # Botões padrão do app (tk.Button)
        btns = tk.Frame(dlg, bg='#ffffff')
        btns.pack(fill=tk.X, padx=12, pady=12)
        def _ok():
            try:
                idx = lb.curselection()
                if not idx:
                    return
                _, nome_imp = itens[idx[0]]
                if not nome_imp:
                    messagebox.showwarning("Impressora não configurada", "Esta impressora não está configurada em Configurações.")
                    return
                sel['valor'] = nome_imp
                dlg.destroy()
            except Exception:
                dlg.destroy()
        def _cancel():
            dlg.destroy()
        tk.Button(btns, text="Imprimir", bg="#17a2b8", fg="white", relief=tk.FLAT, cursor='hand2', command=_ok).pack(side=tk.RIGHT, padx=6)
        tk.Button(btns, text="Cancelar", bg="#6c757d", fg="white", relief=tk.FLAT, cursor='hand2', command=_cancel).pack(side=tk.RIGHT)

        dlg.wait_window()
        return sel['valor']

    # ------------------------- Corretor ortográfico -------------------------
    def _habilitar_corretor_ortografico(self):
        """Configura tags e binds para correção ortográfica em PT.
        Requer 'pyspellchecker'. Se ausente, ignora silenciosamente.
        Também ativa o menu de contexto de sugestões do BaseModule."""
        try:
            txt = getattr(self, 'editor_texto', None)
            if not txt:
                return
            if SpellChecker is None:
                # Dependência ausente: não ativa
                return
            # Inicializa dicionário em português
            if self._spell is None:
                try:
                    self._spell = SpellChecker(language='pt')
                except Exception:
                    self._spell = None
                    return
            # Configura a tag visual das palavras incorretas
            try:
                txt.tag_configure('misspelled', underline=True, foreground='#c82333')
            except Exception:
                pass

            # Bind com debounce para não pesar a digitação
            def _on_key_release(event=None):
                # cancela análise pendente
                if self._spell_after_id:
                    try:
                        self.frame.after_cancel(self._spell_after_id)
                    except Exception:
                        pass
                    self._spell_after_id = None
                # agenda nova verificação
                def _run():
                    self._verificar_ortografia()
                self._spell_after_id = self.frame.after(250, _run)

            txt.bind('<KeyRelease>', _on_key_release, add=True)
            # Primeira verificação do conteúdo inicial (após prefill)
            self._verificar_ortografia()
            # Menu de contexto com sugestões (usa utilitário do BaseModule)
            try:
                self._bind_spell_menu(txt, language='pt', limit_chars=30000)
            except Exception:
                pass
        except Exception:
            pass

    def _verificar_ortografia(self):
        """Varre o texto e marca palavras desconhecidas pelo dicionário PT."""
        if not self._spell or not getattr(self, 'editor_texto', None):
            return
        txt = self.editor_texto
        try:
            # Limita análise para não travar em textos muito grandes
            conteudo = txt.get('1.0', 'end-1c')
            limite = 30000
            analisado = conteudo[:limite]

            # Limpa marcações anteriores
            txt.tag_remove('misspelled', '1.0', tk.END)

            # Encontra palavras (inclui letras acentuadas)
            for m in re.finditer(r"\b[\wÀ-ÖØ-öø-ÿ]+\b", analisado, flags=re.UNICODE):
                palavra = m.group(0)
                # Ignora números puros e tokens muito curtos
                if palavra.isdigit() or len(palavra) <= 2:
                    continue
                # Checa no dicionário
                try:
                    if palavra.lower() in self._spell:  # acesso rápido ao dicionário
                        continue
                except Exception:
                    pass
                # Método padrão da lib
                try:
                    if palavra.lower() not in self._spell:  # redundância segura
                        # Confere probabilidade: unknown retorna conjunto de não reconhecidas
                        if palavra.lower() in self._spell.unknown([palavra.lower()]):
                            start = f"1.0+{m.start()}c"
                            end = f"1.0+{m.end()}c"
                            txt.tag_add('misspelled', start, end)
                except Exception:
                    # Em caso de erro, não marca
                    continue
        except Exception:
            pass

    def _preencher_cabecalho_novo_prontuario(self):
        """Insere cabeçalho com nome do paciente e data atual no editor, se estiver vazio."""
        try:
            if not getattr(self, 'paciente_selecionado', None):
                return
            txt = getattr(self, 'editor_texto', None)
            if not txt:
                return
            # Não sobrescreve se já houver conteúdo
            if txt.get("1.0", "end-1c").strip():
                return
            nome_paciente = self.paciente_selecionado.get('nome', 'Paciente')
            data_hoje = datetime.now().strftime('%d/%m/%Y')
            header = f"Paciente: {nome_paciente}\nData: {data_hoje}\n\n"
            txt.insert("1.0", header)
        except Exception:
            pass

    def _inserir_assinatura_medico(self):
        """Insere duas linhas em branco e, na sequência, o nome do médico e o CRM no final do texto."""
        try:
            txt = getattr(self, 'editor_texto', None) or getattr(self, 'texto_prontuario', None)
            if not txt:
                return
            # Obtém dados do médico a partir do usuário logado via banco (FK usuario_id) com fallback por nome
            medico_info = self._buscar_medico_por_usuario()
            usr = getattr(self, 'usuario_dict', None) or {}
            # Nome exibido: sempre prioriza o nome do médico (tabela medicos); nunca usar login
            nome = (medico_info.get('nome') if medico_info else None) or (usr.get('nome') or '')
            nome = (nome or 'Médico').strip()
            crm = (medico_info.get('crm') if medico_info else None) or (usr.get('crm') or usr.get('CRM') or usr.get('imp') or usr.get('codigo_medico') or '')
            crm = (crm or '').strip()
            try:
                print(f"[DEBUG][Prontuario] assinatura_medico: nome='{nome}', crm='{crm}', medico_info_resolvido={bool(medico_info)}")
            except Exception:
                pass
            assinatura = f"{nome}\nCRM: {crm}" if crm else f"{nome}\nCRM: ______"
            # Garante separação de duas linhas em branco antes da assinatura
            # e centraliza horizontalmente o bloco inserido
            try:
                txt.tag_configure('center', justify='center')
            except Exception:
                pass
            start_idx = txt.index(tk.END)
            txt.insert(tk.END, "\n\n" + assinatura)
            end_idx = txt.index(tk.END)
            try:
                txt.tag_add('center', start_idx, end_idx)
            except Exception:
                pass
            try:
                # Reaplica verificação ortográfica, se ativo
                self._verificar_ortografia()
            except Exception:
                pass
        except Exception:
            pass

    def _buscar_medico_por_usuario(self):
        """Busca nome e CRM do médico vinculado ao usuário logado.
        Estratégia:
        1) Tenta SELECT por medicos.usuario_id = usuarios.id (se a coluna existir).
        2) Se não achar, tenta vincular automaticamente por (nome + telefone) e repete (1).
        3) Fallback: tenta por nome e filtra por telefone (quando disponível), senão usa o primeiro.
        Retorna dict {nome, crm} ou None.
        """
        try:
            usr = getattr(self, 'usuario_dict', None) or {}
            user_id = usr.get('id')
            if not user_id:
                return None
            db = self._obter_conexao_valida()
            if not db:
                return None
            cur = db.cursor()
            try:
                print(f"[DEBUG][Prontuario] buscar_medico_por_usuario: user_id={user_id}")
            except Exception:
                pass
            # 1) por usuario_id
            try:
                cur.execute("SELECT nome, crm FROM medicos WHERE usuario_id = %s LIMIT 1", (user_id,))
                row = cur.fetchone()
                if row:
                    try:
                        print(f"[DEBUG][Prontuario] médico encontrado por usuario_id: nome='{row[0]}', crm='{row[1]}'")
                    except Exception:
                        pass
                    return {"nome": row[0], "crm": row[1]}
            except Exception:
                # coluna pode não existir; segue para fallback
                pass
            # 2) por nome
            nome_usr = (usr.get('nome') or '').strip()
            if nome_usr:
                try:
                    cur.execute("SELECT id, nome, crm, telefone, usuario_id FROM medicos WHERE LOWER(nome) = LOWER(%s) LIMIT 1", (nome_usr,))
                    row = cur.fetchone()
                    if row:
                        try:
                            print(f"[DEBUG][Prontuario] médico encontrado por nome: nome='{row[1]}', crm='{row[2]}'")
                        except Exception:
                            pass
                        return {"nome": row[1], "crm": row[2]}
                except Exception:
                    pass
            # 2) Fallback por nome (pode gerar múltiplos; pega o primeiro)
            try:
                nome_usr = (usr.get('nome') or '').strip()
                tel_usr = (usr.get('telefone') or '').strip()
                if not nome_usr:
                    return None
                cur.execute("SELECT id, nome, crm, telefone, usuario_id FROM medicos WHERE LOWER(nome) = LOWER(%s)", (nome_usr,))
                rows = cur.fetchall() or []
                if not rows:
                    try:
                        print(f"[DEBUG][Prontuario] nenhum médico encontrado por nome='{nome_usr}'")
                    except Exception:
                        pass
                    return None
                # Se houver telefone do usuário, tente match exato de telefone normalizado
                if tel_usr:
                    nt_usr = self._normalizar_telefone(tel_usr)
                    candidatos = []
                    for r in rows:
                        nt_med = self._normalizar_telefone(r[3] or '')
                        if nt_med and nt_med == nt_usr:
                            candidatos.append(r)
                    if len(candidatos) == 1:
                        r = candidatos[0]
                        try:
                            print(f"[DEBUG][Prontuario] médico por nome+telefone: nome='{r[1]}', crm='{r[2]}'")
                        except Exception:
                            pass
                        return {"nome": r[1], "crm": r[2]}
                    # se múltiplos, preferir aquele que já tem usuario_id preenchido
                    for r in candidatos:
                        if r[4]:
                            try:
                                print(f"[DEBUG][Prontuario] múltiplos por nome+telefone, escolhendo com usuario_id: nome='{r[1]}', crm='{r[2]}'")
                            except Exception:
                                pass
                            return {"nome": r[1], "crm": r[2]}
                else:
                    # Sem telefone: só vincula se houver exatamente um médico com o nome e ainda sem usuario_id
                    sem_vinculo = [r for r in rows if not r[4]]
                    if len(sem_vinculo) != 1:
                        try:
                            print(f"[DEBUG][Prontuario] vinculação por nome (sem telefone) ambígua: candidatos_sem_vinculo={len(sem_vinculo)}")
                        except Exception:
                            pass
                        return None
                    r = sem_vinculo[0]
                    try:
                        print(f"[DEBUG][Prontuario] único médico por nome (sem telefone): nome='{r[1]}', crm='{r[2]}'")
                    except Exception:
                        pass
                    return {"nome": r[1], "crm": r[2]}
                # fallback: primeiro registro por nome
                r = rows[0]
                try:
                    print(f"[DEBUG][Prontuario] fallback por nome: nome='{r[1]}', crm='{r[2]}'")
                except Exception:
                    pass
                return {"nome": r[1], "crm": r[2]}
            except Exception:
                return None
        except Exception:
            return None

    def _normalizar_telefone(self, tel: str) -> str:
        """Mantém apenas dígitos para comparação (ex.: '(11) 99999-0000' -> '11999990000')."""
        try:
            return re.sub(r"\D+", "", tel or "")
        except Exception:
            return tel or ""

    def _tentar_vincular_medico_por_nome_telefone(self):
        """Se encontrar exatamente um médico com mesmo nome e telefone do usuário e sem vinculação,
        atualiza medicos.usuario_id.
        """
        try:
            usr = getattr(self, 'usuario_dict', None) or {}
            user_id = usr.get('id')
            nome_usr = (usr.get('nome') or '').strip()
            tel_usr = (usr.get('telefone') or '').strip()
            if not (user_id and nome_usr):
                try:
                    print("[DEBUG][Prontuario] vinculação ignorada: faltam user_id/nome")
                except Exception:
                    pass
                return False
            db = self._obter_conexao_valida()
            if not db:
                try:
                    print("[DEBUG][Prontuario] vinculação ignorada: sem conexão DB válida")
                except Exception:
                    pass
                return False
            cur = db.cursor()
            nt_usr = self._normalizar_telefone(tel_usr)
            # Busca todos com mesmo nome (case-insensitive)
            cur.execute("SELECT id, nome, crm, telefone, usuario_id FROM medicos WHERE LOWER(nome) = LOWER(%s)", (nome_usr,))
            rows = cur.fetchall() or []
            # Filtra por telefone normalizado igual
            candidatos = []
            if tel_usr:
                for r in rows:
                    nt_med = self._normalizar_telefone(r[3] or '')
                    if nt_med and nt_med == nt_usr:
                        candidatos.append(r)
                if len(candidatos) != 1:
                    try:
                        print(f"[DEBUG][Prontuario] vinculação: candidatos por nome+telefone = {len(candidatos)} (esperado 1)")
                    except Exception:
                        pass
                    return False
                r = candidatos[0]
            else:
                # Sem telefone: só vincula se houver exatamente um médico com o nome e ainda sem usuario_id
                sem_vinculo = [r for r in rows if not r[4]]
                if len(sem_vinculo) != 1:
                    try:
                        print(f"[DEBUG][Prontuario] vinculação por nome (sem telefone) ambígua: candidatos_sem_vinculo={len(sem_vinculo)}")
                    except Exception:
                        pass
                    return False
                r = sem_vinculo[0]
            if r[4]:
                # já vinculado
                try:
                    print(f"[DEBUG][Prontuario] vinculação: já possuía usuario_id (medico_id={r[0]})")
                except Exception:
                    pass
                return False
            # Atualiza vinculação
            cur.execute("UPDATE medicos SET usuario_id = %s WHERE id = %s", (user_id, r[0]))
            try:
                db.commit()
            except Exception:
                pass
            try:
                print(f"[DEBUG][Prontuario] vinculação realizada: medico_id={r[0]}, user_id={user_id}")
            except Exception:
                pass
            return True
        except Exception:
            return False

    def _mostrar_prontuario_modal(self, prontuario):
        """Exibe um prontuário histórico em uma janela modal de visualização."""
        dlg = tk.Toplevel(self.frame)
        dlg.title(f"Prontuário #{prontuario.get('id', '')}")
        dlg.geometry("700x500")
        dlg.transient(self.frame)
        dlg.grab_set()
        body = tk.Frame(dlg, bg='#f0f2f5', padx=10, pady=10)
        body.pack(fill=tk.BOTH, expand=True)

        # Conteúdo
        txt = scrolledtext.ScrolledText(body, wrap=tk.WORD, font=('Arial', 10), state=tk.NORMAL)
        txt.pack(fill=tk.BOTH, expand=True)
        txt.insert(tk.END, prontuario.get('conteudo', ''))
        txt.config(state=tk.DISABLED)

        # Fechar
        tk.Button(body, text="Fechar", command=dlg.destroy).pack(pady=8)
    
    def _salvar_novo_prontuario(self):
        """Salva um novo prontuário."""
        # Valida paciente selecionado
        if not self.paciente_selecionado:
            messagebox.showwarning("Aviso", "Selecione um paciente antes de salvar o prontuário.")
            return
        
        # Garante que há um usuário logado (usuario_id) para atribuir ao prontuário
        usuario_id = None
        try:
            usuario_id = (getattr(self, 'usuario_dict', None) or {}).get('id')
        except Exception:
            usuario_id = None
        if not usuario_id:
            messagebox.showwarning(
                "Aviso",
                "Não foi possível identificar o usuário logado. Faça login novamente."
            )
            return
        
        # Obtém o conteúdo do editor exatamente como está (sem alterações)
        origem = getattr(self, 'editor_texto', None) or getattr(self, 'texto_prontuario', None)
        conteudo = origem.get(1.0, tk.END).strip() if origem else ''

        if not conteudo:
            messagebox.showinfo("Aviso", "O conteúdo do prontuário não pode estar vazio")
            return
        
        # Monta os dados conforme schema da tabela 'prontuarios'
        primeira_linha = (conteudo.splitlines()[0] if conteudo.splitlines() else "").strip()
        titulo = (primeira_linha or "Prontuário")[0:255]

        # Se veio da agenda, captura o ID da consulta para vincular e atualizar status
        consulta_id_ctx = None
        try:
            consulta_id_ctx = getattr(self.controller, '_preselect_consulta_id', None)
            if consulta_id_ctx is not None:
                try:
                    consulta_id_ctx = int(consulta_id_ctx)
                except Exception:
                    pass
        except Exception:
            consulta_id_ctx = None

        dados = {
            "paciente_id": self.paciente_selecionado["id"],
            "usuario_id": usuario_id,
            "conteudo": conteudo,
            "data": datetime.now().strftime('%Y-%m-%d'),  # campo 'data' é DATE na tabela
            "titulo": titulo,
            "consulta_id": consulta_id_ctx if consulta_id_ctx else None  # manter None se não houver consulta vinculada
        }
        
        sucesso, resultado = self.prontuario_controller.criar_prontuario(dados)
        
        if sucesso:
   
            # Se houver consulta vinculada, marca como 'Atendido' na agenda
            if consulta_id_ctx:
                try:
                    conn = self._obter_conexao_valida()
                    agenda_ctrl = AgendaController(conn)
                    ok, msg = agenda_ctrl.atualizar_status_consulta(int(consulta_id_ctx), 'Atendido')
                    if not ok:
                        try:
                            print(f"[DEBUG][Prontuario] Falha ao atualizar status da consulta {consulta_id_ctx}: {msg}")
                        except Exception:
                            pass
                    # Sinaliza para a agenda que há atualização pendente de UI
                    try:
                        setattr(self.controller, '_agenda_needs_refresh', True)
                    except Exception:
                        pass
                except Exception as e:
                    try:
                        print(f"[DEBUG][Prontuario] Erro ao atualizar status da consulta {consulta_id_ctx}: {e}")
                    except Exception:
                        pass
                # Limpa o contexto para evitar reaproveitamento indevido
                try:
                    setattr(self.controller, '_preselect_consulta_id', None)
                except Exception:
                    pass

            # Recarrega o painel de prontuários
            self._carregar_prontuarios()
        else:
            messagebox.showerror("Erro", f"Não foi possível criar o prontuário: {resultado}")
    
    def _on_change_busca(self, event=None):
        """Dispara busca com pequeno atraso ao digitar, como em Consultas."""
        # cancela agendamento anterior
        if hasattr(self, '_busca_after_id') and self._busca_after_id:
            try:
                self.frame.after_cancel(self._busca_after_id)
            except Exception:
                pass
            self._busca_after_id = None
        # agenda nova execução
        def _run():
            termo = self.busca_var.get().strip()
            if termo:
                self._buscar_paciente()
            else:
                # limpar lista se vazio
                for item in self.resultados_tree.get_children():
                    self.resultados_tree.delete(item)
        self._busca_after_id = self.frame.after(200, _run)

    def _selecionar_modelo(self):
        """Abre uma janela para selecionar um modelo de texto para inserir no prontuário."""
        # Verifica se está em modo de edição
        if not self.modo_edicao:
            messagebox.showinfo("Aviso", "É necessário estar em modo de edição para inserir modelos")
            return
        
        # Verificação de permissão para modelos:
        # Aceita se: nivel='medico' OU houver medico_id mapeado OU houver usuario_id (id do usuário logado)
        nivel_usr = ((self.usuario_dict or {}).get('nivel') or '').strip().lower()
        usuario_id = (self.usuario_dict or {}).get('id')
        if not (nivel_usr == 'medico' or self.medico_id or usuario_id):
            messagebox.showinfo("Aviso", "É necessário estar logado para usar modelos")
            return
        if nivel_usr == 'medico' and not self.medico_id:
            try:
                print("[DEBUG][Prontuario] prosseguindo com inserção de modelo: nivel=medico porém medico_id não mapeado; usaremos fallback de identificadores")
            except Exception:
                pass
        
        # Busca os modelos do médico com fallback de identificadores
        try:
            print("[DEBUG][Prontuario] _selecionar_modelo: iniciando listagem de modelos")
        except Exception:
            pass
        try:
            modelos = self._listar_modelos_compat()
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao listar modelos de texto: {e}")
            return
        
        if not modelos:
            messagebox.showinfo(
                "Nenhum Modelo",
                "Você ainda não possui modelos de texto disponíveis."
            )
            return
        
        # Cria um diálogo modal
        dialogo = tk.Toplevel(self.frame)
        dialogo.title("Selecionar Modelo de Texto")
        dialogo.geometry("900x650")
        dialogo.transient(self.frame)
        dialogo.grab_set()
        dialogo.focus_set()
        dialogo.configure(bg='#f0f2f5')
        
        # Frame principal
        frame_principal = tk.Frame(dialogo, bg='#f0f2f5', padx=15, pady=15)
        frame_principal.pack(fill=tk.BOTH, expand=True)
        
        # Título
        tk.Label(
            frame_principal,
            text="Selecione um Modelo para Inserir",
            font=('Arial', 14, 'bold'),
            bg='#f0f2f5',
            fg='#333333'
        ).pack(fill=tk.X, pady=(0, 15))
        
        # Frame para a lista de modelos
        frame_lista = tk.Frame(frame_principal, bg='#f0f2f5')
        frame_lista.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        # Lista de modelos com scrollbar
        frame_listbox = tk.Frame(frame_lista, bg='#f0f2f5')
        frame_listbox.pack(fill=tk.BOTH, expand=True)
        
        # Listbox com os modelos (altura maior)
        listbox = tk.Listbox(
            frame_listbox,
            font=('Arial', 11),
            selectbackground='#4a6fa5',
            selectforeground='white',
            activestyle='none',
            height=20
        )
        listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Scrollbar
        scrollbar = ttk.Scrollbar(
            frame_listbox,
            orient="vertical",
            command=listbox.yview
        )
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        listbox.configure(yscrollcommand=scrollbar.set)
        
        # Preenche a listbox com os modelos
        modelo_ids = []
        for modelo in modelos:
            listbox.insert(tk.END, modelo["nome"])
            modelo_ids.append(modelo["id"])
        
        # Removido preview do modelo por solicitação
        
        # Frame para os botões
        frame_botoes = tk.Frame(frame_principal, bg='#f0f2f5')
        frame_botoes.pack(fill=tk.X, pady=(0, 5))
        
        # Função para inserir o modelo selecionado
        def inserir_modelo():
            selecao = listbox.curselection()
            if not selecao:
                messagebox.showinfo("Aviso", "Selecione um modelo para inserir")
                return
            
            modelo_id = modelo_ids[selecao[0]]
            modelo = self.prontuario_controller.buscar_modelo_por_id(modelo_id)
            
            if not modelo:
                return
            
            # Obtém o conteúdo do modelo exatamente como foi salvo
            conteudo = modelo["conteudo"]
            # Se o conteúdo possuir diretiva <<font:N>>, aplica no editor para manter padrão visual
            try:
                m = re.search(r"^\s*<<font:(\d{1,2})>>\s*$", conteudo, flags=re.IGNORECASE | re.MULTILINE)
                if m:
                    n = int(m.group(1))
                    if 6 <= n <= 12:
                        try:
                            self.editor_texto.configure(font=('Arial', n))
                        except Exception:
                            pass
            except Exception:
                pass
            
            # Insere no NOVO editor do prontuário (somente a tela atual)
            try:
                existente = self.editor_texto.get("1.0", "end-1c").strip()
                if existente:
                    self.editor_texto.insert(tk.END, "\n\n")
            except Exception:
                pass
            # Remove quaisquer linhas de diretiva de fonte do conteúdo antes de inserir
            try:
                linhas = conteudo.replace('\r\n', '\n').replace('\r', '\n').split('\n')
                linhas = [ln for ln in linhas if not (ln.strip().lower().startswith('<<font:') and ln.strip().endswith('>>'))]
                # Remove cabeçalho (Paciente/Médico/Data) do INÍCIO do conteúdo do modelo, se existir
                def _strip_initial_header(ls):
                    i = 0
                    seen = {'paciente': False, 'medico': False, 'data': False}
                    while i < len(ls):
                        s = (ls[i] or '').strip()
                        if not s:
                            i += 1
                            continue
                        low = s.lower()
                        if low.startswith('paciente:') or low.startswith('paciente :'):
                            seen['paciente'] = True; i += 1; continue
                        if low.startswith('médico:') or low.startswith('medico:') or low.startswith('médico :') or low.startswith('medico :'):
                            seen['medico'] = True; i += 1; continue
                        if low.startswith('data:') or low.startswith('data :'):
                            seen['data'] = True; i += 1; continue
                        break
                    if seen['paciente'] or seen['medico'] or seen['data']:
                        while i < len(ls) and not (ls[i] or '').strip():
                            i += 1
                        return ls[i:]
                    return ls
                linhas = _strip_initial_header(linhas)
                conteudo_limp = "\n".join(linhas)
            except Exception:
                conteudo_limp = conteudo
            self.editor_texto.insert(tk.END, conteudo_limp)
            # Assinatura do médico duas linhas abaixo do modelo inserido
            try:
                self._inserir_assinatura_medico()
            except Exception:
                pass
            
            # Fecha o diálogo
            dialogo.destroy()
        
        # Duplo clique para inserir diretamente
        listbox.bind("<Double-1>", lambda e: inserir_modelo())
        
        # Botões de ação
        tk.Button(
            frame_botoes,
            text="Inserir",
            command=inserir_modelo,
            bg="#4a6fa5",
            fg="white",
            relief=tk.FLAT,
            cursor='hand2'
        ).pack(side=tk.RIGHT, padx=5)

        tk.Button(
            frame_botoes,
            text="Cancelar",
            command=dialogo.destroy,
            bg="#6c757d",
            fg="white",
            relief=tk.FLAT,
            cursor='hand2'
        ).pack(side=tk.RIGHT, padx=5)

        # Removido botão 'Gerenciar Modelos' por solicitação
        
        # Centraliza o diálogo
        dialogo.update_idletasks()
        width = dialogo.winfo_width()
        height = dialogo.winfo_height()
        x = (dialogo.winfo_screenwidth() // 2) - (width // 2)
        y = (dialogo.winfo_screenheight() // 2) - (height // 2)
        dialogo.geometry(f"{width}x{height}+{x}+{y}")
        
        # Seleciona o primeiro item da lista, se houver
        if listbox.size() > 0:
            listbox.selection_set(0)

    def _selecionar_receita(self):
        """Abre um diálogo para selecionar uma Receita cadastrada e inseri-la no novo prontuário.
        Após inserir, adiciona a assinatura do médico (Nome e CRM) abaixo, como no padrão atual.
        """
        # Garantir edição ativa
        if not self.modo_edicao:
            messagebox.showinfo("Aviso", "É necessário estar em modo de edição para inserir receitas")
            return
        # Garantir um CadastroController válido (do controller principal ou on-demand)
        cad = getattr(self, 'cadastro_controller', None)
        if not cad:
            try:
                self.cadastro_controller = CadastroController()
                # alinhar conexão de banco na instância criada
                conn = self._obter_conexao_valida()
                if conn and hasattr(self.cadastro_controller, 'db'):
                    self.cadastro_controller.db = conn
            except Exception:
                messagebox.showerror("Erro", "Módulo de cadastro não disponível para listar receitas.")
                return
        
        # Resolver medico_id para filtrar receitas
        medico_id = self._resolver_medico_id_para_receitas()
        if not medico_id:
            messagebox.showinfo("Aviso", "Não foi possível identificar o médico logado para listar receitas.")
            return
        
        # Listar receitas do médico
        try:
            receitas = self.cadastro_controller.listar_receitas_por_medico(medico_id) or []
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao listar receitas: {e}")
            return
        if not receitas:
            messagebox.showinfo("Receitas", "Nenhuma receita cadastrada para este médico.")
            return
        
        # Montar diálogo
        dlg = tk.Toplevel(self.frame)
        dlg.title("Selecionar Receita")
        dlg.transient(self.frame)
        dlg.grab_set()
        dlg.configure(bg='#f0f2f5')
        
        container = tk.Frame(dlg, bg='#f0f2f5', padx=15, pady=15)
        container.pack(fill=tk.BOTH, expand=True)
        tk.Label(container, text="Selecione uma Receita para Inserir", font=('Arial', 14, 'bold'), bg='#f0f2f5', fg='#333').pack(fill=tk.X, pady=(0, 12))
        
        list_frame = tk.Frame(container, bg='#f0f2f5')
        list_frame.pack(fill=tk.BOTH, expand=True)
        listbox = tk.Listbox(list_frame, font=('Arial', 11), selectbackground='#4a6fa5', selectforeground='white', activestyle='none', height=20)
        listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb = ttk.Scrollbar(list_frame, orient='vertical', command=listbox.yview)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        listbox.configure(yscrollcommand=sb.set)
        
        receita_ids = []
        for r in receitas:
            nome = r.get('nome') or f"Receita {r.get('id')}"
            listbox.insert(tk.END, nome)
            receita_ids.append(r.get('id'))
        
        def inserir_receita():
            idxs = listbox.curselection()
            if not idxs:
                messagebox.showinfo("Aviso", "Selecione uma receita para inserir")
                return
            rid = receita_ids[idxs[0]]
            try:
                rec = self.cadastro_controller.obter_receita_por_id(rid)
            except Exception as e:
                messagebox.showerror("Erro", f"Erro ao obter receita: {e}")
                return
            if not rec:
                return
            texto = (rec.get('texto') or '').strip()
            try:
                existente = self.editor_texto.get("1.0", "end-1c").strip()
                if existente:
                    self.editor_texto.insert(tk.END, "\n\n")
            except Exception:
                pass
            if texto:
                # Remove quaisquer linhas de diretiva de fonte antes de inserir a receita
                try:
                    linhas_r = texto.replace('\r\n', '\n').replace('\r', '\n').split('\n')
                    linhas_r = [ln for ln in linhas_r if not (ln.strip().lower().startswith('<<font:') and ln.strip().endswith('>>'))]
                    # Remove cabeçalho (Paciente/Médico/Data) do INÍCIO do texto da receita, se existir
                    def _strip_initial_header(ls):
                        i = 0
                        seen = {'paciente': False, 'medico': False, 'data': False}
                        while i < len(ls):
                            s = (ls[i] or '').strip()
                            if not s:
                                i += 1
                                continue
                            low = s.lower()
                            if low.startswith('paciente:') or low.startswith('paciente :'):
                                seen['paciente'] = True; i += 1; continue
                            if low.startswith('médico:') or low.startswith('medico:') or low.startswith('médico :') or low.startswith('medico :'):
                                seen['medico'] = True; i += 1; continue
                            if low.startswith('data:') or low.startswith('data :'):
                                seen['data'] = True; i += 1; continue
                            break
                        if seen['paciente'] or seen['medico'] or seen['data']:
                            while i < len(ls) and not (ls[i] or '').strip():
                                i += 1
                            return ls[i:]
                        return ls
                    linhas_r = _strip_initial_header(linhas_r)
                    texto_limp = "\n".join(linhas_r)
                except Exception:
                    texto_limp = texto
                self.editor_texto.insert(tk.END, texto_limp)
            # Assinatura do médico abaixo
            try:
                self._inserir_assinatura_medico()
            except Exception:
                pass
            dlg.destroy()
        
        # Ações
        btns = tk.Frame(container, bg='#f0f2f5')
        btns.pack(fill=tk.X, pady=(8, 0))
        tk.Button(btns, text="Inserir", command=inserir_receita, bg="#4a6fa5", fg="white", relief=tk.FLAT, cursor='hand2').pack(side=tk.RIGHT, padx=5)
        tk.Button(btns, text="Cancelar", command=dlg.destroy, bg="#6c757d", fg="white", relief=tk.FLAT, cursor='hand2').pack(side=tk.RIGHT, padx=5)
        
        # Centralizar
        try:
            dlg.update_idletasks()
            w, h = dlg.winfo_width(), dlg.winfo_height()
            sw, sh = dlg.winfo_screenwidth(), dlg.winfo_screenheight()
            x = (sw - w) // 2
            y = (sh - h) // 2
            dlg.geometry(f"{w}x{h}+{x}+{y}")
        except Exception:
            pass
        
        if listbox.size() > 0:
            listbox.selection_set(0)

    def _resolver_medico_id_para_receitas(self):
        """Tenta resolver o medico_id a partir do usuário logado para filtrar receitas.
        Estratégia: procurar na tabela medicos por usuario_id; fallback por nome.
        """
        try:
            conn = self._obter_conexao_valida()
            if not conn:
                return None
            usr = getattr(self, 'usuario_dict', None) or {}
            user_id = usr.get('id')
            nome_usr = (usr.get('nome') or '').strip()
            cur = conn.cursor()
            # 1) por usuario_id
            if user_id:
                try:
                    cur.execute("SELECT id FROM medicos WHERE usuario_id = %s LIMIT 1", (user_id,))
                    r = cur.fetchone()
                    if r:
                        return r[0]
                except Exception:
                    pass
            # 2) por nome
            if nome_usr:
                try:
                    cur.execute("SELECT id FROM medicos WHERE LOWER(nome) = LOWER(%s) LIMIT 1", (nome_usr,))
                    r = cur.fetchone()
                    if r:
                        return r[0]
                except Exception:
                    pass
        except Exception:
            return None
        return None

    def _gerenciar_modelos(self):
        """Abre uma janela para gerenciar modelos de texto."""
        # Verifica permissão: permite quando houver medico_id OU usuario_id (compatível com modelos por usuário)
        if not (self.medico_id or (self.usuario_dict or {}).get('id')):
            messagebox.showwarning("Acesso Restrito", "É necessário estar logado para gerenciar modelos de texto.")
            return

        # Cria uma janela de diálogo
        dialogo = tk.Toplevel(self.frame)
        dialogo.title("Gerenciar Modelos de Texto")
        dialogo.geometry("800x600")
        dialogo.transient(self.frame)
        dialogo.grab_set()
        dialogo.configure(bg='#f0f2f5')

        # Frame principal
        frame_principal = tk.Frame(dialogo, bg='#f0f2f5', padx=15, pady=15)
        frame_principal.pack(fill=tk.BOTH, expand=True)

        # Título
        tk.Label(
            frame_principal,
            text="Gerenciar Modelos de Texto",
            font=('Arial', 16, 'bold'),
            bg='#f0f2f5',
            fg='#333333'
        ).pack(fill=tk.X, pady=(0, 15))

        # Descrição
        tk.Label(
            frame_principal,
            text="Aqui você pode criar, editar e excluir seus modelos de texto personalizados.",
            font=('Arial', 11),
            bg='#f0f2f5',
            fg='#555555',
            wraplength=750
        ).pack(fill=tk.X, pady=(0, 15))

        # Frame para a lista e o conteúdo
        frame_conteudo = tk.Frame(frame_principal, bg='#f0f2f5')
        frame_conteudo.pack(fill=tk.BOTH, expand=True)

        # Frame para a lista de modelos
        frame_lista = tk.Frame(frame_conteudo, bg='#f0f2f5', width=250)
        frame_lista.pack(side=tk.LEFT, fill=tk.BOTH, padx=(0, 10))

        # Título da lista
        tk.Label(
            frame_lista,
            text="Seus Modelos",
            font=('Arial', 12, 'bold'),
            bg='#f0f2f5',
            fg='#333333'
        ).pack(fill=tk.X, pady=(0, 5))

        # Lista de modelos com scrollbar
        frame_listbox = tk.Frame(frame_lista, bg='#f0f2f5')
        frame_listbox.pack(fill=tk.BOTH, expand=True)

        listbox = tk.Listbox(
            frame_listbox,
            font=('Arial', 11),
            selectbackground='#4a6fa5',
            selectforeground='white',
            activestyle='none',
            height=15
        )
        listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        scrollbar = ttk.Scrollbar(
            frame_listbox,
            orient="vertical",
            command=listbox.yview
        )
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        listbox.configure(yscrollcommand=scrollbar.set)

        # Botões para a lista (sem criação de modelos nesta tela)
        frame_botoes_lista = tk.Frame(frame_lista, bg='#f0f2f5')
        frame_botoes_lista.pack(fill=tk.X, pady=(10, 0))

        # Frame para o conteúdo do modelo
        frame_conteudo_direito = tk.Frame(frame_conteudo, bg='#f0f2f5')
        frame_conteudo_direito.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

        # Mensagem inicial
        tk.Label(
            frame_conteudo_direito,
            text="Selecione um modelo para visualizar ou editar",
            font=('Arial', 12),
            bg='#f0f2f5',
            fg='#555555'
        ).pack(pady=50)

        # Armazena os IDs dos modelos
        modelo_ids = []

        # Função para carregar os modelos
        def carregar_modelos():
            # Limpa a lista
            listbox.delete(0, tk.END)
            modelo_ids.clear()

            # Carrega os modelos do médico com fallback de identificadores
            try:
                modelos = self._listar_modelos_compat()
            except Exception as e:
                messagebox.showerror("Erro", f"Erro ao listar modelos de texto: {e}")
                return

            if not modelos:
                listbox.insert(tk.END, "Nenhum modelo encontrado")
                return

            # Adiciona os modelos à lista
            for modelo in modelos:
                listbox.insert(tk.END, modelo['nome'])
                modelo_ids.append(modelo['id'])

        # Função para exibir o modelo selecionado
        def exibir_modelo_selecionado(event):
            # Limpa o frame de conteúdo
            for widget in frame_conteudo_direito.winfo_children():
                widget.destroy()

            # Obtém o índice selecionado
            try:
                indice = listbox.curselection()[0]
            except IndexError:
                return

            # Verifica se há modelos
            if not modelo_ids:
                return

            # Obtém o ID do modelo
            modelo_id = modelo_ids[indice]

            # Busca o modelo
            modelo = self.prontuario_controller.buscar_modelo_texto_por_id(modelo_id)

            if not modelo:
                return

            # Frame para o nome do modelo
            frame_nome = tk.Frame(frame_conteudo_direito, bg='#f0f2f5')
            frame_nome.pack(fill=tk.X, pady=(0, 10))

            tk.Label(
                frame_nome,
                text="Nome:",
                font=('Arial', 11),
                bg='#f0f2f5'
            ).pack(side=tk.LEFT)

            nome_var = tk.StringVar(value=modelo['nome'])
            nome_entry = ttk.Entry(
                frame_nome,
                textvariable=nome_var,
                font=('Arial', 11),
                width=40
            )
            nome_entry.pack(side=tk.LEFT, padx=(5, 0))

            # Frame para o conteúdo do modelo
            frame_texto = tk.Frame(frame_conteudo_direito, bg='#f0f2f5')
            frame_texto.pack(fill=tk.BOTH, expand=True, pady=(0, 10))

            tk.Label(
                frame_texto,
                text="Conteúdo:",
                font=('Arial', 11),
                bg='#f0f2f5'
            ).pack(anchor="w")

            # Área de texto para o conteúdo
            texto_modelo = scrolledtext.ScrolledText(
                frame_texto,
                wrap=tk.WORD,
                font=('Arial', 11),
                height=15
            )
            texto_modelo.pack(fill=tk.BOTH, expand=True)
            texto_modelo.insert(tk.END, modelo['conteudo'])

            # Frame para os botões
            frame_botoes = tk.Frame(frame_conteudo_direito, bg='#f0f2f5')
            frame_botoes.pack(fill=tk.X)

            # Botão para salvar alterações
            ttk.Button(
                frame_botoes,
                text="Salvar Alterações",
                command=lambda: salvar_alteracoes(modelo_id, nome_var.get(), texto_modelo.get("1.0", tk.END)),
                style='Accent.TButton'
            ).pack(side=tk.RIGHT)

        # Função para salvar alterações no modelo
        def salvar_alteracoes(modelo_id, nome, conteudo):
            # Valida os campos
            if not nome or not conteudo.strip():
                messagebox.showwarning("Campos Obrigatórios", "Preencha todos os campos.")
                return

            # Atualiza o modelo
            sucesso = self.prontuario_controller.atualizar_modelo_texto(
                modelo_id,
                nome,
                conteudo.strip()
            )

            if sucesso:
                messagebox.showinfo("Sucesso", "Modelo atualizado com sucesso!")
                carregar_modelos()
            else:
                messagebox.showerror("Erro", "Não foi possível atualizar o modelo.")

        # Vincula a seleção de um modelo
        listbox.bind("<<ListboxSelect>>", exibir_modelo_selecionado)

        # Carrega os modelos iniciais
        carregar_modelos()

        # Centraliza o diálogo
        dialogo.update_idletasks()
        width = dialogo.winfo_width()
        height = dialogo.winfo_height()
        x = (dialogo.winfo_screenwidth() // 2) - (width // 2)
        y = (dialogo.winfo_screenheight() // 2) - (height // 2)
        dialogo.geometry(f"{width}x{height}+{x}+{y}")

    def _novo_modelo_texto(self, frame_conteudo, listbox, modelo_ids):
        """Cria um novo modelo de texto."""
        # Limpa o frame de conteúdo
        for widget in frame_conteudo.winfo_children():
            widget.destroy()

        # Título
        tk.Label(
            frame_conteudo,
            text="Novo Modelo",
            font=('Arial', 12, 'bold'),
            bg='#f0f2f5',
            fg='#333333'
        ).pack(fill=tk.X, pady=(0, 10))

        # Frame para o nome do modelo
        frame_nome = tk.Frame(frame_conteudo, bg='#f0f2f5')
        frame_nome.pack(fill=tk.X, pady=(0, 10))

        tk.Label(
            frame_nome,
            text="Nome:",
            font=('Arial', 11),
            bg='#f0f2f5'
        ).pack(side=tk.LEFT)

        nome_var = tk.StringVar()
        nome_entry = ttk.Entry(
            frame_nome,
            textvariable=nome_var,
            font=('Arial', 11),
            width=40
        )
        nome_entry.pack(side=tk.LEFT, padx=(5, 0))

        # Frame para o conteúdo
        frame_texto = tk.Frame(frame_conteudo, bg='#f0f2f5')
        frame_texto.pack(fill=tk.BOTH, expand=True, pady=(0, 10))

        tk.Label(
            frame_texto,
            text="Conteúdo:",
            font=('Arial', 11),
            bg='#f0f2f5'
        ).pack(anchor="w")

        # Área de texto para o conteúdo
        texto = scrolledtext.ScrolledText(
            frame_texto,
            wrap=tk.WORD,
            font=('Arial', 11),
            height=15
        )
        texto.pack(fill=tk.BOTH, expand=True)

        # Frame para os botões
        frame_botoes = tk.Frame(frame_conteudo, bg='#f0f2f5')
        frame_botoes.pack(fill=tk.X, pady=(0, 5))

        # Função para salvar o novo modelo
        def salvar_modelo():
            nome = nome_var.get().strip()
            conteudo = texto.get(1.0, tk.END).strip()

            if not nome or not conteudo:
                messagebox.showwarning("Campos Obrigatórios", "Preencha todos os campos.")
                return

            # Cria o novo modelo
            dados = {
                "nome": nome,
                "conteudo": conteudo,
                # modelos_texto agora usa usuario_id
                "usuario_id": (self.usuario_dict or {}).get('id')
            }

            sucesso, resultado = self.prontuario_controller.criar_modelo_texto(dados)

            if sucesso:
                messagebox.showinfo("Sucesso", "Modelo criado com sucesso!")

                # Recarrega a lista de modelos
                listbox.delete(0, tk.END)
                modelo_ids.clear()

                modelos = self.prontuario_controller.listar_modelos_texto((self.usuario_dict or {}).get('id'))
                for modelo in modelos:
                    listbox.insert(tk.END, modelo['nome'])
                    modelo_ids.append(modelo['id'])

                # Limpa o formulário
                nome_var.set("")
                texto.delete(1.0, tk.END)
            else:
                messagebox.showerror("Erro", f"Não foi possível criar o modelo: {resultado}")

        # Botões de ação
        ttk.Button(
            frame_botoes,
            text="Salvar",
            command=salvar_modelo,
            style='Accent.TButton'
        ).pack(side=tk.RIGHT, padx=5)

        ttk.Button(
            frame_botoes,
            text="Cancelar",
            command=lambda: frame_conteudo.winfo_children()[0].destroy()
        ).pack(side=tk.RIGHT, padx=5)

        # Define o foco no campo de nome
        nome_entry.focus_set()

    def exibir(self):
        """Exibe o módulo sem adicionar elementos extras na tela."""
        # Apenas garante que o frame do módulo esteja na tela
        self.frame.pack(fill=tk.BOTH, expand=True)

    def _criar_novo_prontuario(self):
        """Prepara o editor à direita para um novo prontuário, limpo e pronto para edição."""
        # Garante que o container da direita exista
        if not hasattr(self, 'prontuario_content_frame') or not self.prontuario_content_frame.winfo_exists():
            return
        # Limpa o conteúdo atual do painel direito
        for w in self.prontuario_content_frame.winfo_children():
            w.destroy()
        # Renderiza novamente o editor de novo prontuário
        self._render_editor_novo_prontuario(self.prontuario_content_frame)
        # Foco no editor
        try:
            self.editor_texto.focus_set()
        except Exception:
            pass
